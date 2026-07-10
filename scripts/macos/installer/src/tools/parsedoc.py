#!/usr/bin/env python3
# parsedoc <文件> — 提取文档文字/内容到 stdout（Linux 版）。
# PDF(pymupdf 正文 + pdfplumber 表格；无文字层→Qwen-VL OCR) / docx(python-docx) /
# rtf,odt,doc(pandoc) / xlsx(openpyxl) / 纯文本。CC 直接调，无需选库。
import sys, os, tempfile, subprocess, logging

# 静音 pdfminer/pdfplumber 对字体子集的 WARNING（如 chromium 生成的 PDF 的
# "Could not get FontBBox from font descriptor"）——纯噪声，不影响文字/表格提取。
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

HERE = os.path.dirname(os.path.abspath(__file__))

def die(m): sys.stderr.write(m + "\n"); sys.exit(1)

if len(sys.argv) < 2 or not os.path.isfile(sys.argv[1]):
    die("用法: parsedoc <文件>")
F = sys.argv[1]
ext = F.rsplit(".", 1)[-1].lower() if "." in os.path.basename(F) else ""

def parse_pdf(path):
    import fitz
    doc = fitz.open(path)
    pages = [(i, p.get_text().strip()) for i, p in enumerate(doc)]
    total = "".join(t for _, t in pages).strip()
    if len(total) >= 20:
        for i, t in pages:
            print("=== 第%d页 ===" % (i + 1)); print(t)
        tables = []
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for pi, pg in enumerate(pdf.pages):
                    for tbl in (pg.extract_tables() or []): tables.append((pi, tbl))
        except Exception as e:
            sys.stderr.write("[parsedoc] pdfplumber: %s\n" % e)
        if not tables:
            try:
                for pi, pg in enumerate(doc):
                    for t in pg.find_tables().tables: tables.append((pi, t.extract()))
            except Exception as e:
                sys.stderr.write("[parsedoc] find_tables: %s\n" % e)
        for n, (pi, tbl) in enumerate(tables):
            print("\n=== 第%d页·表格%d ===" % (pi + 1, n + 1))
            for row in tbl:
                print(" | ".join("" if c is None else str(c).replace("\n", " ").strip() for c in row))
    else:
        sys.stderr.write("[parsedoc] PDF 无文字层，转图走 Qwen-VL OCR…\n")
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            fd, tmp = tempfile.mkstemp(suffix=".png"); os.close(fd); pix.save(tmp)
            try:
                r = subprocess.run([sys.executable, os.path.join(HERE, "describe-image.py"), tmp,
                    "逐字读出这一页所有文字，保留顺序与表格结构。"],
                    capture_output=True, text=True, timeout=150)
                print("=== 第%d页(OCR) ===" % (i + 1)); print((r.stdout or r.stderr).strip())
            finally:
                os.remove(tmp)

def parse_docx(path):
    import docx
    d = docx.Document(path)
    for p in d.paragraphs:
        if p.text.strip(): print(p.text)
    for k, tb in enumerate(d.tables):
        print("\n=== 表格%d ===" % (k + 1))
        for row in tb.rows:
            print(" | ".join(c.text.strip() for c in row.cells))

def parse_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    for ws in wb.worksheets:
        print("=== 工作表:%s ===" % ws.title)
        for r in ws.iter_rows(values_only=True):
            print(" | ".join("" if v is None else str(v) for v in r))

def parse_via_pandoc(path):  # rtf/odt/doc：用 pandoc 转纯文本
    try:
        r = subprocess.run(["pandoc", path, "-t", "plain"], capture_output=True, text=True, timeout=120)
        if r.returncode == 0:
            sys.stdout.write(r.stdout)
        else:
            die("[parsedoc] pandoc 失败: " + (r.stderr or "")[:300])
    except FileNotFoundError:
        die("[parsedoc] 需要 pandoc 才能解析 ." + ext)

if ext == "pdf":
    parse_pdf(F)
elif ext == "docx":
    parse_docx(F)
elif ext in ("rtf", "odt", "doc"):
    parse_via_pandoc(F)
elif ext == "xlsx":
    parse_xlsx(F)
elif ext in ("txt", "md", "csv", "tsv", "json", "log", "html", "xml", "yaml", "yml"):
    sys.stdout.write(open(F, encoding="utf-8", errors="replace").read())
else:
    die("[parsedoc] 不支持: ." + ext)
